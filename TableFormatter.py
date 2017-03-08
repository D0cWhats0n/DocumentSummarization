# -*- coding: utf-8 -*-
"""
Created on Tue Mar 07 11:12:11 2017

@author: JohannesM
"""
from docx import Document
import numpy as np
import pandas as pd

def readTableAsDataFrame(path):
    document = Document(path)
    table_df = pd.DataFrame()
    for table in document.tables:
        for index,column in enumerate(table.columns):
            column_name = "column_" + str(index)
            paragraph_np = np.empty([len(column.cells)], dtype=object)
            #print len(column.cells)
            for cell_index,cell in enumerate(column.cells):
                paragraph_string = ''
                for paragraph in cell.paragraphs:
                    paragraph_string += paragraph.text.encode('utf-8')
                paragraph_np[cell_index] = paragraph_string        
            table_df[column_name] = pd.Series(paragraph_np) 
    return table_df;  

def readFileAsString(path):
    if file.endswith(".txt"):
        with open(path, 'r') as myfile:
            data=myfile.read().replace('\n', '')
        return "\n" + data
    elif file.endswith(".docx"):
        document = Document(path)
        fullText = ""
        for para in document.paragraphs:
            fullText += "\n" + para.text
        for table in document.tables:
            for column in table.columns:
                for cell in column.cells:
                    for paragraph in cell.paragraphs:
                       fullText += "\n" + paragraph.text    
        return '\n' + fullText  
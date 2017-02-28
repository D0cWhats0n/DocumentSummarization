# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""
from docx import Document
from sumy.parsers.html import HtmlParser
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.sum_basic import SumBasicSummarizer as Summarizer
from sumy.nlp.stemmers import Stemmer
from sumy.utils import get_stop_words
import os

#set language and sentence count
LANGUAGE = "german"
SENTENCES_COUNT = 20
filePath = os.path.dirname(__file__)+ "/textFiles/"

concat_string = ""

def readFileAsString(path):
    if file.endswith(".txt"):
        with open(path, 'r') as myfile:
            data=myfile.read().replace('\n', '')
        return "\n" + data
    elif file.endswith(".docx"):
        document = Document(path)
        fullText = ""
        for para in document.paragraphs:
            fullText += para.text
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                       fullText += "\n" + paragraph.text    
        return '\n' + fullText 


#read files from directory
for file in os.listdir(filePath):
    concat_string += readFileAsString(filePath + file) 
    
        
parser = PlaintextParser.from_string(concat_string, Tokenizer(LANGUAGE))
stemmer = Stemmer(LANGUAGE)

summarizer = Summarizer(stemmer)
summarizer.stop_words = get_stop_words(LANGUAGE)

for sentence in summarizer(parser.document, SENTENCES_COUNT):
    print(sentence)




